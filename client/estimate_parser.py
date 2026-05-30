# client/estimate_parser.py
# 견적서 파싱 - pdfplumber/openpyxl/python-docx 기반
# WEB 복구판: PDF/XLSX 다중 견적서, 업체별 품목, WISENIC류 가공견적서 보정

import re, os, tempfile, shutil, subprocess, sys
from pathlib import Path

ITEM_KEYS  = ["품목명", "품명", "품명및규격", "품목명/규격", "품명/규격",
              "품명 및 규격", "내역", "명칭", "부품명", "item", "description",
              "model", "modle", "모델", "모델명", "제품명", "자재명", "부품", "part name", "item description",
              "品名", "品 名", "DESCRIPTION"]
SPEC_KEYS  = ["규격", "사양", "형번", "도번", "품번", "품 번", "part no", "partno", "p/n", "spec",
              "規格", "規 格", "SPECIFICATION"]
QTY_KEYS   = ["수량", "qty", "q'ty", "quantity", "數量", "數 量", "数量", "QUANTITY"]
PRICE_KEYS = ["단가", "단위가격", "공급단가", "unit price", "u/price", "price",
              "單價", "單 價", "単価", "UNIT PRICE"]
AMT_KEYS   = ["금액", "공급가액", "공급금액", "amount", "합계금액", "합계",
              "金額", "金 額", "金額", "AMOUNT"]
UNIT_KEYS  = ["단위", "unit", "單位", "単位"]
MAKER_KEYS = ["비고", "비고(제조사)", "비고/제조사", "제조사", "메이커", "maker", "brand", "업체", "업체명", "거래처", "공급사", "supplier", "재질", "표면처리", "재질&표면처리", "재질및표면처리",
              "備考", "備 考", "REMARK"]

SKIP_WORDS = ["소계", "합계", "합 계", "총합계", "이하여백", "이 하 여 백", "이하 여백",
              "total", "subtotal", "특기사항", "납기일자", "유효일자", "가능납기",
              "결제조건", "납품장소", "비고란", "발주시", "확인 부탁", "견적금액",
              "견적내용", "견적서", "아래와 같이", "하기와 같이", "귀중", "귀하",
              "전화번호", "사업자번호", "대표자", "주소", "fax", "공급자", "수요자"]

def _s(v):
    return "" if v is None else str(v).strip()

def _n(s):
    return re.sub(r"\s+", "", str(s or "")).lower()

def _match(cell, keys):
    c = _n(cell)
    if not c:
        return False
    return any(_n(k) in c or c in _n(k) for k in keys)

def _num(s):
    s = re.sub(r"\s+", "", str(s or ""))
    # 날짜는 숫자로 오인하지 않도록 제외
    if re.match(r"^\d{4}[-/.년]\d{1,2}", s):
        return 0.0
    s = re.sub(r"[^\d.]", "", s)
    try:
        return float(s)
    except Exception:
        return 0.0

def _has_num(s):
    return bool(re.search(r"\d", str(s or "")))

def _skip(name):
    n = _n(name)
    if not n:
        return True
    return any(_n(w) in n for w in SKIP_WORDS)



def _is_skip_row_text(text: str) -> bool:
    """품목 행이 아닌 합계/조건/공급자/주소 행을 더 엄격하게 걸러낸다."""
    n = _n(text)
    if not n:
        return True
    hard = [
        '소계', '합계', '총합계', 'subtotal', 'total', 'vat', '부가세', '세액', '견적금액', '총견적',
        '납기', '납품', '결제', '결재', '유효', '조건', '비고란', '특기사항', '이하여백', '이하여백',
        '사업자등록', '사업자번호', '대표자', '주소', '전화', '팩스', 'fax', 'tel', '공급자', '수요자',
        '은행', '계좌', '예금주', '담당자', '작성자', '견적서', 'quote', 'quotation'
    ]
    return any(_n(k) in n for k in hard)


def _is_no_col_header(cell: str) -> bool:
    c = _n(cell)
    return c in {'no', 'no.', '번호', '순번', '연번', '번호no'} or c.startswith('no.')


def _is_money_like(cell: str) -> bool:
    raw = str(cell or '').strip()
    if not raw:
        return False
    if re.search(r"\d{4}[-/.년]\d{1,2}", raw):
        return False
    v = _num(raw)
    if v <= 0:
        return False
    # 수량 1~99와 구분하기 위해 콤마/원/₩ 또는 3자리 이상을 금액 후보로 본다.
    return bool(re.search(r"[,₩원]|\d{3,}", raw))


def _is_qty_like(cell: str) -> bool:
    raw = str(cell or '').strip()
    if not raw:
        return False
    if re.search(r"\d{4}[-/.년]\d{1,2}", raw):
        return False
    v = _num(raw)
    return 0 < v <= 99999 and not re.search(r"[,₩원]", raw)


def _clean_item_text(value: str) -> str:
    text = re.sub(r"\s+", " ", str(value or '')).strip(' :;/\t')
    # 셀 안에 품명: 값처럼 들어온 경우 라벨 제거
    text = re.sub(r"^(품명|품목명|제품명|명칭|description|item)\s*[:：]\s*", "", text, flags=re.I).strip()
    return text


def _clean_spec_text(value: str) -> str:
    text = re.sub(r"\s+", " ", str(value or '')).strip(' :;/\t')
    text = re.sub(r"^(규격|사양|형번|품번|도번|spec|p/n|part\s*no)\s*[:：]\s*", "", text, flags=re.I).strip()
    return text


def _estimate_row_amount_consistency(item: dict) -> bool:
    try:
        qty = float(item.get('quantity') or 0)
        price = float(item.get('unit_price') or 0)
        amt = float(item.get('amount') or 0)
    except Exception:
        return False
    if qty <= 0:
        return False
    if price > 0 and amt > 0:
        expected = qty * price
        # 견적서에 부가세/할인이 섞인 경우를 감안해 어느 정도 오차 허용
        return abs(expected - amt) <= max(5, abs(expected) * 0.08)
    return price > 0 or amt > 0


def _normalize_and_filter_items(items):
    """파서별 결과를 공통 후처리한다.

    - 합계/부가세/납기/주소 행 제거
    - 단가/금액/수량 보정
    - 중복 제거
    - 비고/제조사 자동 채움 제거
    """
    out, seen = [], set()
    for src in items or []:
        it = dict(src or {})
        name = _clean_item_text(it.get('item_name') or it.get('name') or it.get('item') or '')
        spec = _clean_spec_text(it.get('spec') or it.get('standard') or it.get('part_no') or '')
        if spec == name:
            spec = ''
        row_text = ' '.join(str(it.get(k, '') or '') for k in ('item_name', 'spec', 'maker', 'note'))
        if not _looks_item_name(name) or _is_skip_row_text(row_text):
            continue
        try:
            qty = float(it.get('quantity') or it.get('qty') or 0)
        except Exception:
            qty = 0
        if qty <= 0:
            qty = 1
        try:
            price = float(it.get('unit_price') or it.get('price') or 0)
        except Exception:
            price = 0
        try:
            amt = float(it.get('amount') or it.get('amt') or 0)
        except Exception:
            amt = 0
        if price <= 0 and amt > 0 and qty > 0:
            price = amt / qty
        if amt <= 0 and price > 0 and qty > 0:
            amt = price * qty
        # 금액 정보가 전혀 없는 줄은 품목 행으로 보기 어렵다.
        if price <= 0 and amt <= 0:
            continue
        unit = str(it.get('unit') or 'EA').strip().upper()
        unit_map = {'EACH': 'EA', 'PCS': 'EA', 'PC': 'EA', '개': 'EA'}
        unit = unit_map.get(unit, unit or 'EA')
        if unit not in ['EA', 'M', 'KG', 'L', 'BOX', 'SET', 'ROLL', '식', '본', '매']:
            unit = 'EA'
        norm = {
            'item_name': name,
            'spec': spec,
            'unit_price': int(round(price)),
            'quantity': int(qty) if float(qty).is_integer() else qty,
            'unit': unit,
            'maker': '',
            'note': '',
            'amount': int(round(amt)),
        }
        if it.get('material_code'):
            norm['material_code'] = it.get('material_code')
        if it.get('item_group'):
            norm['item_group'] = it.get('item_group')
        key = (_n(norm['item_name']), _n(norm['spec']), str(norm['quantity']), str(norm['amount']))
        if key in seen:
            continue
        seen.add(key)
        out.append(norm)
    return out


def _parse_amount_anchored_rows(rows):
    """헤더가 없거나 헤더가 깨진 표에서 금액 열을 기준으로 품목 행을 추정한다.

    업체별 견적서 양식 차이가 큰 경우에도 보통 우측에 `수량/단가/금액`이 있으므로,
    우측 금액 후보 2개와 그 앞의 수량 후보를 찾고 왼쪽 텍스트를 품명/규격으로 나눈다.
    """
    out = []
    pending_name = ''
    pending_spec = ''
    for raw_row in rows or []:
        cells = [_s(c) for c in (raw_row or [])]
        cells = [c for c in cells if c]
        if len(cells) < 4:
            # 단독 품명/규격 줄은 다음 금액 행과 연결하기 위해 저장
            txt = ' '.join(cells).strip()
            if txt and not _is_skip_row_text(txt):
                if _looks_part_no(txt):
                    pending_spec = txt
                elif _looks_item_name(txt):
                    pending_name = txt
            continue
        joined = ' '.join(cells)
        if _is_skip_row_text(joined) or _is_header_like_line(joined):
            continue
        money_idx = [i for i, c in enumerate(cells) if _is_money_like(c)]
        if not money_idx:
            continue
        # 일반적으로 마지막 금액=amount, 그 앞 금액=unit_price
        amount_i = money_idx[-1]
        price_i = money_idx[-2] if len(money_idx) >= 2 else None
        qty_i = None
        # 단가 앞쪽에서 수량 후보를 찾는다. No. 열은 제외하기 위해 뒤에서부터 탐색한다.
        limit = price_i if price_i is not None else amount_i
        for i in range(limit - 1, -1, -1):
            if _is_qty_like(cells[i]):
                qty_i = i
                break
        if qty_i is None:
            continue
        qty = max(1, int(_num(cells[qty_i])))
        amount = int(_num(cells[amount_i]))
        price = int(_num(cells[price_i])) if price_i is not None else (amount // qty if qty else 0)
        if price <= 0 and amount and qty:
            price = amount // qty
        # 품명/규격 후보는 수량 앞쪽 텍스트에서 선택한다.
        left = [c for i, c in enumerate(cells[:qty_i]) if not _is_no_col_header(c) and not c.replace('.', '', 1).isdigit()]
        if not left and (pending_name or pending_spec):
            left = [x for x in [pending_name, pending_spec] if x]
        if not left:
            continue
        part_candidates = [c for c in left if _looks_part_no(c)]
        spec = _clean_spec_text(part_candidates[-1] if part_candidates else pending_spec)
        name_candidates = [c for c in left if c != spec and _looks_item_name(c) and not _is_skip_row_text(c)]
        if pending_name and not name_candidates:
            name_candidates.append(pending_name)
        # 너무 긴 문장형 셀보다 마지막 명확한 텍스트 후보를 선호
        name = _clean_item_text(name_candidates[-1] if name_candidates else '')
        if not name:
            continue
        out.append({
            'item_name': name,
            'spec': spec,
            'unit_price': price,
            'quantity': qty,
            'unit': 'EA',
            'maker': '',
            'note': '',
            'amount': amount or price * qty,
        })
        pending_name = ''
        pending_spec = ''
    return _normalize_and_filter_items(out)


def _infer_item_group_from_text(text: str) -> str:
    """견적서 본문/파일명에서 품목군을 추정해 DB 보조값으로만 저장한다."""
    t = _n(text)
    if not t:
        return ""
    # 더 구체적인 분류를 우선한다.
    rules = [
        ("판금품", ["판금", "sheetmetal", "sheetmetal", "sheet-metal", "sheet metal"]),
        ("선반품", ["선반", "turning", "lathe", "cnc선반"]),
        ("밀링품", ["밀링", "milling", "mct"]),
        ("가공품", ["가공", "machining", "machinepart", "가공견적"]),
        ("구매품", ["구매품", "구매", "purchase"]),
        ("전장품", ["전장", "전장품", "전장부품", "전장부품류", "cable", "케이블"]),
        ("프레임", ["프레임", "frame"]),
        ("석정반", ["석정반", "granite"]),
    ]
    for group, keys in rules:
        for k in keys:
            if _n(k) in t:
                return group
    return ""


def _enrich_items_hidden_fields(items, context: str = ""):
    """화면 입력값은 유지하되 DB/품질관리용 숨은 보조값을 채운다.

    - material_code: 견적서의 규격/도번(spec)을 기본 자재코드로 저장
    - item_group: 판금품/가공품/선반품 등 문맥에서 추정한 품목군 저장
    """
    group = _infer_item_group_from_text(context)
    for it in items or []:
        if not it.get("material_code") and it.get("spec"):
            it["material_code"] = it.get("spec") or ""
        if not it.get("item_group") and group:
            it["item_group"] = group
    return items

def _looks_part_no(s: str) -> bool:
    s = _s(s)
    return bool(re.search(r"[A-Za-z]", s) and re.search(r"\d", s) and len(s) >= 5)

def _looks_item_name(s: str) -> bool:
    s = _s(s)
    if not s or _skip(s):
        return False
    if re.match(r"^\d{4}[-/.년]", s):
        return False
    return any(ch.isalnum() for ch in s)

_VENDOR_JUNK_SUFFIXES = [
    "아래와 같이", "다음과 같이", "와 같이", "과 같이", "하기와 같이",
    "견적드립니다", "견적 드립니다", "견적서를", "견적서", " 귀중", " 귀하",
]

def _clean_vendor_name(name: str) -> str:
    name = re.sub(r"\s+", " ", str(name or "")).strip(" :-_/\t")
    for junk in _VENDOR_JUNK_SUFFIXES:
        if junk in name:
            name = name[:name.index(junk)].strip()
    # 같은 행에 납품장소/주소 문구와 업체명이 붙는 구형 Excel 보정
    # 예: "귀사지정도 두성자동화㈜" -> "두성자동화㈜"
    tokens = [t.strip() for t in name.split() if t.strip()]
    for i, tok in enumerate(tokens):
        if tok == "주식회사" and i + 1 < len(tokens):
            # `주식회사 제이씨 모션`처럼 회사명이 두 단어인 경우까지 보존하되,
            # `상 호 주식회사 동신정밀 성 명 ...`의 `성` 같은 다음 필드명은 제외한다.
            stop_tokens = {"성", "성명", "명", "대표", "대표자", "주소", "주", "소", "업", "태", "품", "목"}
            tail = []
            for cand in tokens[i + 1:i + 3]:
                if cand in stop_tokens:
                    break
                tail.append(cand)
            name = (tok + " " + " ".join(tail)).strip()
            break
        if tok in ("(주)", "㈜") and i + 1 < len(tokens):
            name = (tok + " " + tokens[i + 1]).strip()
            break
        if ("㈜" in tok or "(주)" in tok or "주식회사" in tok) and len(tok) >= 3:
            name = tok
    # 대표/주소/전화 등 뒤쪽 설명 제거
    for stop in ["대표이사", "대표", "사업자", "주소", "TEL", "Tel", "전화", "FAX", "Fax"]:
        if stop in name:
            name = name.split(stop)[0].strip()
    return name.strip()

def _vendor_from_text(text):
    text = str(text or "")
    # 수신처(이노로보틱스 귀중)보다 공급자/상호 표기를 우선 사용
    patterns = [
        r"공\s*급\s*자\s*[:：]?\s*([가-힣A-Za-z0-9㈜()\s._-]{2,30})",
        r"공\s*급\s*사\s*[:：]?\s*([가-힣A-Za-z0-9㈜()\s._-]{2,30})",
        r"상\s*호\s*[:：]?\s*([가-힣A-Za-z0-9㈜()\s._-]{2,30})",
        r"회\s*사\s*명\s*[:：]?\s*([가-힣A-Za-z0-9㈜()\s._-]{2,30})",
        r"회사명/대표\s*[:：]?\s*([가-힣A-Za-z0-9㈜()\s._-]{2,30})",
        r"취\s*급\s*품\s*목\s*[:：]?\s*([가-힣A-Za-z0-9㈜()\s._-]{2,30})",
        r"([가-힣A-Za-z0-9㈜()\s._-]{2,30}(?:주식회사|㈜|\(주\)))",
    ]
    for pattern in patterns:
        for m in re.finditer(pattern, text, re.MULTILINE):
            name = _clean_vendor_name(m.group(1))
            if name and "이노로보틱스" not in name and len(name) >= 2:
                # 과하게 긴 주소/업태 문자열 방지
                for stop in ["사업자", "대표", "주소", "업태", "전화", "fax", "Fax"]:
                    if stop in name:
                        name = name.split(stop)[0].strip()
                return name
    return ""

def _vendor_from_rows(rows):
    # 행 단위로 공급자/상호 옆 셀 탐색
    for row in rows[:30]:
        cells = [_s(c) for c in row]
        for i, c in enumerate(cells):
            nc = _n(c)
            if nc in ["공급자", "공급사", "상호", "회사명"] or "공급자" == nc:
                for nxt in cells[i+1:i+4]:
                    name = _clean_vendor_name(nxt)
                    if name and "이노로보틱스" not in name and not any(k in name for k in ["사업자", "대표자", "주소"]):
                        return name
    return _vendor_from_text("\n".join("\t".join(_s(c) for c in row) for row in rows[:30]))

def _find_header(rows):
    """견적서 표 헤더 위치/컬럼 추정.

    단일 헤더행뿐 아니라 병합/2단 헤더를 합친 후보도 검사한다.
    반환 col에는 `_header_span`이 포함될 수 있다.
    """
    def is_model_header(cell):
        c = _n(cell)
        return c in ["model", "modle", "모델", "모델명", "제품명", "자재명", "부품", "partname", "itemdescription"] or "model" in c or "modle" in c

    def is_name_header(cell):
        c = _n(cell)
        return any(k in c for k in ["품명", "품목명", "품목", "명칭", "내역", "description", "item", "품명규격", "품명및규격"])

    def is_size_spec_header(cell):
        c = _n(cell)
        return any(k in c for k in ["규격", "사양", "형번", "spec", "standard"]) and not any(k in c for k in ["품번", "도번"])

    def is_part_no_header(cell):
        c = _n(cell)
        return any(k in c for k in ["품번", "품 번", "도번", "partno", "part no", "p/n", "품목코드", "자재코드", "materialcode"])

    def combine_rows(row_a, row_b=None, row_c=None):
        rows2 = [row_a]
        if row_b is not None:
            rows2.append(row_b)
        if row_c is not None:
            rows2.append(row_c)
        max_len = max((len(r) for r in rows2 if r), default=0)
        out = []
        for j in range(max_len):
            parts = []
            for r in rows2:
                if r is not None and j < len(r) and _s(r[j]):
                    parts.append(_s(r[j]))
            out.append(' '.join(parts).strip())
        return out

    def score_mapping(m):
        score = 0
        if 'item' in m: score += 4
        if 'spec' in m: score += 2
        if 'qty' in m: score += 3
        if 'price' in m: score += 3
        if 'amt' in m: score += 3
        if 'unit' in m: score += 1
        return score

    best = (-1, {}, -1)
    for i in range(min(len(rows), 80)):
        variants = [(rows[i], 1)]
        if i + 1 < len(rows):
            variants.append((combine_rows(rows[i], rows[i+1]), 2))
        if i + 2 < len(rows):
            variants.append((combine_rows(rows[i], rows[i+1], rows[i+2]), 3))
        for row, span in variants:
            m = {}
            model_idx = name_idx = size_spec_idx = part_no_idx = None
            no_idx = None
            for j, cell in enumerate(row):
                if cell is None:
                    continue
                c = _n(cell)
                if not c:
                    continue
                if no_idx is None and _is_no_col_header(cell):
                    no_idx = j
                if model_idx is None and is_model_header(cell):
                    model_idx = j
                if name_idx is None and is_name_header(cell):
                    name_idx = j
                if size_spec_idx is None and is_size_spec_header(cell):
                    size_spec_idx = j
                if part_no_idx is None and is_part_no_header(cell):
                    part_no_idx = j
                if _match(cell, QTY_KEYS) and 'qty' not in m:
                    m['qty'] = j
                if _match(cell, PRICE_KEYS) and 'price' not in m:
                    m['price'] = j
                if _match(cell, AMT_KEYS) and 'amt' not in m:
                    m['amt'] = j
                if _match(cell, UNIT_KEYS) and 'unit' not in m:
                    m['unit'] = j
                if _match(cell, MAKER_KEYS) and 'maker' not in m:
                    m['maker'] = j
            if model_idx is not None:
                m['item'] = model_idx
            elif name_idx is not None:
                m['item'] = name_idx
            elif part_no_idx is not None and size_spec_idx is None:
                # 품번/도번만 있고 품명이 다음 열인 양식 보정
                m['item'] = min(part_no_idx + 1, len(row) - 1)
            if model_idx is not None and part_no_idx is not None:
                m['spec'] = part_no_idx
            elif size_spec_idx is not None:
                m['spec'] = size_spec_idx
            elif part_no_idx is not None:
                m['spec'] = part_no_idx
            if 'item' in m and 'qty' in m and ('price' in m or 'amt' in m):
                m['_header_span'] = span
                # No 열을 item으로 잡은 잘못된 경우를 방지
                if no_idx is not None and m.get('item') == no_idx and part_no_idx is not None:
                    m['item'] = min(part_no_idx + 1, len(row) - 1)
                return i, m
            if score_mapping(m) > score_mapping(best[1]):
                m['_header_span'] = span
                best = (i, m, span)
    # 부족한 헤더라도 item/qty/amount 조합이면 인정
    if best[0] >= 0 and 'item' in best[1] and 'qty' in best[1] and ('price' in best[1] or 'amt' in best[1]):
        return best[0], best[1]
    return -1, {}

def _parse_table(rows):
    if not rows:
        return []
    clean_rows = [[_s(c) for c in (r or [])] for r in rows if r]
    max_cols = max((len(r) for r in clean_rows), default=0)
    if max_cols < 3:
        return []

    header_idx, col = _find_header(clean_rows)
    if header_idx < 0 or "item" not in col:
        return []

    data_start = header_idx + int(col.get("_header_span", 1) or 1)
    # 서브헤더/가공비 하위구분 행 건너뜀
    while data_start < len(clean_rows):
        nxt = clean_rows[data_start]
        sub_kws = ["원재료", "가공", "밀링", "기타", "후처리", "소계", "자재비"]
        if sum(1 for c in nxt if any(_n(k) in _n(c) for k in sub_kws)) >= 2:
            data_start += 1
        else:
            break

    def getcell(field, row, default=""):
        idx = col.get(field)
        if idx is not None and idx < len(row):
            return _s(row[idx])
        return default

    items = []
    for row in clean_rows[data_start:]:
        raw_name = getcell("item", row)
        spec = getcell("spec", row)

        # 헤더가 어긋난 경우: 품번은 spec, MODEL은 품명으로 재보정
        if (not raw_name or _skip(raw_name)) and spec and col.get("spec", 0)+1 < len(row):
            cand = _s(row[col.get("spec") + 1])
            if _looks_item_name(cand):
                raw_name = cand

        if "\n" in raw_name:
            original_raw_name = raw_name
            parts = [x.strip() for x in str(raw_name or '').split("\n") if x.strip()]
            # 강유시스템류: 품목명/규격 칼럼 하나에 `품명\n규격`이 들어가고
            # 수량/단가/금액은 1세트만 있는 경우는 1개 품목으로 병합한다.
            qty_probe = getcell("qty", row)
            price_probe = getcell("price", row)
            amt_probe = getcell("amt", row)
            if len(parts) >= 2 and (_has_num(qty_probe) or _has_num(price_probe) or _has_num(amt_probe)):
                raw_name = parts[0]
                if not spec or spec == original_raw_name or "\n" in str(spec or ''):
                    spec = " ".join(parts[1:]).strip()
            else:
                items.extend(_parse_merged(row, col))
                continue

        if not _looks_item_name(raw_name):
            continue

        # spec이 비어있고 앞/뒤 셀에 도번처럼 보이는 값이 있으면 사용
        if (not spec or spec == raw_name) and col.get("item", 0) > 0:
            prev = _s(row[col["item"]-1])
            if _looks_part_no(prev):
                spec = prev
        if spec == raw_name:
            spec = ""

        # 행 전체에서 도번/품번처럼 보이는 값을 한 번 더 찾는다.
        # 일부 Excel/PDF는 병합셀/빈셀 때문에 header index가 밀려서
        # LDIO-S-A400-M305-1 같은 도번이 단가 칼럼으로 잘못 들어오는 경우가 있음.
        part_candidates = []
        for c in row:
            cc = _s(c)
            if cc and cc != raw_name and _looks_part_no(cc):
                part_candidates.append(cc)
        if (not spec or spec == raw_name or not _looks_part_no(spec)) and part_candidates:
            spec = part_candidates[0]

        qty_raw   = getcell("qty", row)
        price_raw = getcell("price", row)
        amt_raw   = getcell("amt", row)
        unit_raw  = getcell("unit", row, "EA")
        maker     = getcell("maker", row)

        qty   = max(1, int(_num(qty_raw))) if _has_num(qty_raw) else 1
        price = int(_num(price_raw)) if _has_num(price_raw) else 0
        amt   = int(_num(amt_raw)) if _has_num(amt_raw) else 0
        if price == 0 and amt > 0 and qty > 0:
            price = amt // qty
        # 도번/품번 숫자가 단가로 잘못 들어간 경우 보정
        if price > 1_000_000 and amt > 0 and qty > 0:
            price = amt // qty
        # 가끔 예상단가가 날짜/금액 병합으로 터지는 경우 방지
        if price > 10_000_000_000:
            price = amt // qty if amt and qty else 0

        unit = unit_raw.upper() if _n(unit_raw) else "EA"
        if unit not in ["EA", "M", "KG", "L", "BOX", "SET", "ROLL", "식", "개", "본", "매"]:
            unit = "EA"

        items.append({
            "item_name": raw_name,
            "spec": spec,
            "unit_price": price,
            "quantity": qty,
            "unit": unit,
            "maker": maker,
            "amount": amt or (price * qty if price and qty else 0),
        })
    return _normalize_and_filter_items(items)

def _parse_merged(row, col):
    items = []
    def split_col(field):
        idx = col.get(field)
        if idx is not None and idx < len(row):
            return [x.strip() for x in str(row[idx] or "").split("\n") if x.strip()]
        return []
    names = split_col("item")
    specs = split_col("spec")
    qtys = split_col("qty")
    prices = split_col("price")
    amts = split_col("amt")
    for j, name in enumerate(names):
        if not _looks_item_name(name):
            continue
        qty = max(1, int(_num(qtys[j]))) if j < len(qtys) and _has_num(qtys[j]) else 1
        price = int(_num(prices[j])) if j < len(prices) and _has_num(prices[j]) else 0
        amt = int(_num(amts[j])) if j < len(amts) and _has_num(amts[j]) else 0
        if price == 0 and amt > 0 and qty > 0:
            price = amt // qty
        items.append({"item_name": name, "spec": specs[j] if j < len(specs) else "", "unit_price": price, "quantity": qty, "unit": "EA", "maker": "", "amount": amt})
    return _normalize_and_filter_items(items)

def _row_text_cells(row):
    return [_s(c) for c in (row or []) if _s(c)]


def _is_header_like_line(text: str) -> bool:
    n = _n(text)
    if not n:
        return True
    return any(k in n for k in ['no', '번호', '품명', '규격', '수량', '단가', '금액', '합계', '견적', '공급자', '수요자'])




def _parse_numeric_quote_rows(rows):
    """구형 Excel 견적서 보조 인식.

    예: 두성자동화 양식처럼 `NO / 품명 및 규격 / 수량 / 단가 / 금액 / 비고` 구조이지만
    병합/빈 칸 때문에 일반 테이블 인식이 약해지는 경우를 처리한다.
    """
    clean_rows = [[_s(c) for c in (r or [])] for r in rows if r]
    if not clean_rows:
        return []

    header_i = -1
    cols = {}
    for i, row in enumerate(clean_rows[:100]):
        m = {}
        for j, cell in enumerate(row):
            if not _s(cell):
                continue
            n = _n(cell)
            if any(k in n for k in ["품명", "품목", "규격", "내역", "description", "model"]):
                m.setdefault("item", j)
            if _match(cell, QTY_KEYS):
                m.setdefault("qty", j)
            if _match(cell, PRICE_KEYS):
                m.setdefault("price", j)
            if _match(cell, AMT_KEYS):
                m.setdefault("amt", j)
            if _match(cell, UNIT_KEYS):
                m.setdefault("unit", j)
            if _match(cell, MAKER_KEYS):
                m.setdefault("maker", j)
        if "item" in m and "qty" in m and ("price" in m or "amt" in m):
            header_i, cols = i, m
            break
    if header_i < 0:
        return []

    def get(row, key):
        idx = cols.get(key)
        return _s(row[idx]) if idx is not None and idx < len(row) else ""

    items = []
    for row in clean_rows[header_i+1:]:
        text = " ".join(row)
        if any(_n(k) in _n(text) for k in ["합계", "소계", "이하여백", "이하 여백"]):
            continue
        name = get(row, "item")
        if not _looks_item_name(name):
            continue
        qty_raw = get(row, "qty")
        price_raw = get(row, "price")
        amt_raw = get(row, "amt")
        if not (_has_num(qty_raw) or _has_num(price_raw) or _has_num(amt_raw)):
            continue
        qty = max(1, int(_num(qty_raw))) if _has_num(qty_raw) else 1
        price = int(_num(price_raw)) if _has_num(price_raw) else 0
        amt = int(_num(amt_raw)) if _has_num(amt_raw) else 0
        if price == 0 and amt and qty:
            price = amt // qty
        if not amt and price and qty:
            amt = price * qty
        if price == 0 and amt == 0:
            continue
        spec = ""
        for c in row:
            cc = _s(c)
            if cc and cc != name and _looks_part_no(cc) and not cc.replace('.', '', 1).isdigit():
                spec = cc
                break
        unit = get(row, "unit") or "EA"
        maker = get(row, "maker")
        items.append({
            "item_name": name,
            "spec": spec,
            "unit_price": price,
            "quantity": qty,
            "unit": unit,
            "maker": maker,
            "amount": amt or price * qty,
        })
    return _normalize_and_filter_items(items)

def _parse_loose_rows(rows):
    """헤더가 깨지거나 병합셀 때문에 _find_header가 실패한 견적서 보조 인식.

    형식 예:
      No / 품명 / 규격 / 수량 / 단가 / 금액
      품번 MODEL QTY UNIT PRICE AMOUNT
      품명 규격 수량 금액 비고(제조사)
    """
    items = []
    for row in rows or []:
        cells = _row_text_cells(row)
        if len(cells) < 3:
            continue
        joined = ' '.join(cells)
        if _skip(joined) or _is_header_like_line(joined):
            continue
        nums = []
        for i, c in enumerate(cells):
            if _has_num(c):
                val = _num(c)
                if val:
                    nums.append((i, val, c))
        if not nums:
            continue

        # 금액/단가 후보를 먼저 잡고, 그 앞쪽 작은 숫자를 수량으로 추정한다.
        # 기존 small[-1] 방식은 단가 5,470처럼 작은 단가를 수량으로 오인할 수 있었다.
        money = [(i, v, raw) for i, v, raw in nums if _is_money_like(raw) or v >= 1000]
        price = 0
        amount = 0
        price_idx = None
        if money:
            amount = int(money[-1][1])
            if len(money) >= 2:
                price_idx = money[-2][0]
                price = int(money[-2][1])
        first_money_idx = money[0][0] if money else len(cells)
        qty_idx, qty = None, 1
        qty_candidates = [(i, v, raw) for i, v, raw in nums if 1 <= v <= 99999 and i < first_money_idx and not _is_no_col_header(raw)]
        # EA/개/SET 바로 앞 숫자를 최우선 수량으로 사용
        unit_words = {'EA','EACH','PCS','PC','개','SET','BOX','M','KG','L','ROLL','본','매','식'}
        for i, v, raw in reversed(qty_candidates):
            nxt = cells[i+1].upper() if i + 1 < len(cells) else ''
            if nxt in unit_words:
                qty_idx, qty = i, max(1, int(v))
                break
        if qty_idx is None and qty_candidates:
            qty_idx, qty, _ = qty_candidates[-1]
            qty = max(1, int(qty))
        if price == 0 and amount and qty:
            price = int(amount // qty) if amount else 0

        spec = next((c for c in cells if _looks_part_no(c)), '')
        maker = ''
        # 비고/제조사로 보이는 마지막 텍스트 후보
        for c in reversed(cells):
            if c == spec:
                continue
            if _has_num(c):
                continue
            if _looks_item_name(c) and not _skip(c):
                maker = c
                break

        name = ''
        for i, c in enumerate(cells):
            if i == qty_idx or c == spec or _has_num(c):
                continue
            if not _looks_item_name(c):
                continue
            # maker와 같은 마지막 비고가 먼저 잡히지 않게 앞쪽 텍스트를 품명으로 사용
            name = c
            break
        if not name:
            continue
        if maker == name:
            maker = ''
        # spec이 비었으면 품명 직전/직후의 도번 후보를 사용
        if not spec:
            name_idx = cells.index(name) if name in cells else -1
            for c in (cells[:name_idx] + cells[name_idx+1:] if name_idx >= 0 else cells):
                if _looks_part_no(c):
                    spec = c
                    break
        if not price and amount and qty:
            price = int(amount // qty)
        items.append({
            'item_name': name,
            'spec': spec,
            'unit_price': price,
            'quantity': qty,
            'unit': 'EA',
            'maker': maker,
            'note': maker,
            'amount': amount or (price * qty if price and qty else 0),
        })
    # 같은 품명/규격 중복 제거
    out, seen = [], set()
    for it in items:
        key = (_n(it.get('item_name')), _n(it.get('spec')), int(it.get('quantity') or 0), int(it.get('amount') or 0))
        if key in seen:
            continue
        seen.add(key)
        out.append(it)
    return _normalize_and_filter_items(out)


def _is_multilingual_quote_header(line: str) -> bool:
    """한자/영문 혼합 견적서 헤더 인식.

    예: 品 名 規 格 數 量 單 價 金 額 備 考
        DESCRIPTION SPECIFICATION QUANTITY UNIT PRICE AMOUNT REMARK
    """
    n = _n(line)
    if not n:
        return False
    has_eng = ('description' in n and 'specification' in n and 'quantity' in n and ('unitprice' in n or 'price' in n) and 'amount' in n)
    has_cjk = (('品名' in line.replace(' ', '') or '品 名' in line) and
               ('規格' in line.replace(' ', '') or '規 格' in line) and
               (('數量' in line.replace(' ', '') or '數 量' in line) or ('数量' in line.replace(' ', ''))) and
               (('單價' in line.replace(' ', '') or '單 價' in line) or ('単価' in line.replace(' ', ''))) and
               (('金額' in line.replace(' ', '') or '金 額' in line) or ('金額' in line.replace(' ', ''))))
    return has_eng or has_cjk


def _split_spec_qty_price_line(line: str, default_name: str = ''):
    """LEMO/SPH Tech류 PDF처럼 표 선이 없어 텍스트 라인만 추출되는 견적서 보조 파서.

    지원 예:
      CONNECTOR FGG.1B.304.CLAD72 1EA @ 85,000 ₩85,000
      PFG.1B.304.CLLD72 1 EA @ 101,300 ₩101,300
    """
    raw = re.sub(r"\s+", " ", str(line or '')).strip()
    if not raw or _skip(raw):
        return None
    if _is_multilingual_quote_header(raw):
        return None
    if raw.startswith('*') or raw.startswith('▣'):
        return None

    # 수량+단위는 1EA, 1 EA 모두 허용. 단가 앞 @/₩ 선택 허용.
    pat = re.compile(
        r"^(?:(?P<name>[A-Za-z가-힣][A-Za-z가-힣0-9()/&+.,_\- ]{1,60}?)\s+)?"
        r"(?P<spec>[A-Za-z0-9][A-Za-z0-9./_\-]{3,})\s+"
        r"(?P<qty>\d+(?:\.\d+)?)\s*(?P<unit>[A-Za-z가-힣]+)\s*"
        r"@?\s*₩?\s*(?P<price>[0-9,]+)"
        r"(?:\s+₩?\s*(?P<amount>[0-9,]+))?"
        r"(?:\s+(?P<remark>.*))?$",
        re.IGNORECASE,
    )
    m = pat.match(raw)
    if not m:
        return None
    spec = _s(m.group('spec'))
    if not _looks_part_no(spec):
        return None
    name = _s(m.group('name')) or default_name or '품목'
    # CONNECTOR처럼 첫 행에만 품명이 있고 뒤 행에는 규격만 있는 경우 같은 품명을 유지
    qty = max(1, int(_num(m.group('qty'))))
    unit = (_s(m.group('unit')) or 'EA').upper()
    if unit not in ["EA", "M", "KG", "L", "BOX", "SET", "ROLL", "식", "개", "본", "매"]:
        unit = 'EA'
    price = int(_num(m.group('price')))
    amount = int(_num(m.group('amount'))) if _s(m.group('amount')) else price * qty
    remark = _s(m.group('remark'))
    return {
        'item_name': name,
        'spec': spec,
        'unit_price': price,
        'quantity': qty,
        'unit': unit,
        'maker': remark,
        'note': remark,
        'amount': amount,
    }


def _parse_multilingual_quote_text(text: str, vendor: str = ''):
    lines = [re.sub(r"\s+", " ", l).strip() for l in str(text or '').splitlines() if l.strip()]
    if not lines:
        return []

    start = -1
    for i, line in enumerate(lines):
        if _is_multilingual_quote_header(line):
            start = i + 1
    if start < 0:
        # 헤더가 누락되어도 LEMO식 품번/수량/단가 패턴이 있으면 전체 라인에서 시도
        start = 0

    items = []
    current_name = ''
    stop_words = ['TOTAL', 'T O T A L', '견적유효일', '결재조건', '납기 :', '취 급 품 목', 'TEL', 'FAX', '___']
    for line in lines[start:]:
        upper = line.upper()
        if any(sw.upper() in upper for sw in stop_words) or line.startswith('1. 귀사') or line.startswith('2. 의뢰'):
            # 본문 설명 이후는 품목표가 아니므로 중단. 단, 이미 품목을 찾은 뒤에만 중단한다.
            if items:
                break
        parsed = _split_spec_qty_price_line(line, current_name)
        if not parsed:
            continue
        if parsed.get('item_name') and parsed['item_name'] != '품목':
            current_name = parsed['item_name']
        if vendor and not parsed.get('maker'):
            # 비고(제조사)/업체 구분용으로 vendor를 기본 maker/note에 넣는다.
            parsed['maker'] = vendor
            parsed['note'] = vendor
        items.append(parsed)

    # 중복 제거
    out, seen = [], set()
    for it in items:
        key = (_n(it.get('item_name')), _n(it.get('spec')), int(it.get('quantity') or 0), int(it.get('amount') or 0))
        if key in seen:
            continue
        seen.add(key)
        out.append(it)
    return out



def _normalize_split_amounts_in_line(line: str) -> str:
    """OCR/PDF 텍스트에서 `3 44,000`처럼 단가가 쪼개진 값을 `344,000`으로 보정."""
    s = re.sub(r"\s+", " ", str(line or "")).strip()
    # 3 44,000 / 1 0,600 / 5 6,800 같은 금액 분리 보정.
    # 단, `2 422,000`처럼 수량+단가 사이 공백은 병합하면 안 되므로
    # 앞 토큰이 한 자리 숫자이고 뒤 금액의 콤마 앞 숫자가 1~2자리인 경우만 병합한다.
    s = re.sub(r"\b([1-9])\s+(\d{1,2},\d{3})(?=\D|$)", r"\1\2", s)
    return s


def _is_numeric_token(tok: str) -> bool:
    return bool(re.fullmatch(r"[₩￦]?\d[\d,]*(?:\.\d+)?", str(tok or '').strip()))


def _token_num(tok: str) -> float:
    return _num(tok)


def _parse_structured_quote_line(line: str, pending_spec: str = ''):
    """PDF 표가 추출되지 않는 일반 견적서 텍스트 한 줄을 품목으로 인식.

    지원 예:
      1 GTR1205EC5Z-300LF(Raydent,너트역방향) 2 422,000.0 844,000 84,400
      1 SERVO MOTOR SGM7J-04AFA6C 1 EA 3 44,000 344,000
      1 LM GUIDE 4 140,600 562,400 56,240
    """
    raw = _normalize_split_amounts_in_line(line)
    if not raw or _skip(raw) or _is_header_like_line(raw):
        return None
    if not re.match(r"^\s*\d+\s+", raw):
        return None
    tokens = raw.split()
    if len(tokens) < 4 or not re.fullmatch(r"\d+", tokens[0]):
        return None

    units = {"EA", "EACH", "SET", "BOX", "M", "MM", "개", "식", "대", "본", "매"}

    def is_money(tok):
        if not _is_numeric_token(tok):
            return False
        return _token_num(tok) >= 100

    qty_idx = None
    for i in range(1, len(tokens) - 2):
        if not _is_numeric_token(tokens[i]):
            continue
        qv = _token_num(tokens[i])
        if not (1 <= qv <= 9999):
            continue
        j = i + 1
        if j < len(tokens) and tokens[j].upper() in units:
            j += 1
        # 수량 뒤에는 단가/금액이 바로 이어지는 경우만 인정한다.
        if j + 1 < len(tokens) and is_money(tokens[j]) and is_money(tokens[j + 1]):
            qty_idx = i
            break
    if qty_idx is None:
        return None

    qty = max(1, int(_token_num(tokens[qty_idx])))
    price_start = qty_idx + 1
    unit = "EA"
    if price_start < len(tokens) and tokens[price_start].upper() in units:
        unit = "EA" if tokens[price_start].upper() == "EACH" else tokens[price_start].upper()
        price_start += 1
    if price_start + 1 >= len(tokens):
        return None
    price = int(_token_num(tokens[price_start]))
    amount = int(_token_num(tokens[price_start + 1]))
    # 세액/비고가 뒤에 더 붙어도 무시
    name_tokens = tokens[1:qty_idx]
    if not name_tokens:
        return None

    spec = pending_spec or ""
    item_name = " ".join(name_tokens).strip()
    # 마지막 토큰이 도번/품번이면 규격으로 분리
    if len(name_tokens) >= 2 and _looks_part_no(name_tokens[-1]):
        spec = name_tokens[-1]
        item_name = " ".join(name_tokens[:-1]).strip()
    # 품명만 코드형으로 온 경우는 품명에 그대로 두되, 규격은 공란 유지
    if not item_name:
        return None
    if not _looks_item_name(item_name):
        return None
    return {
        "item_name": item_name,
        "spec": spec,
        "unit_price": price,
        "quantity": qty,
        "unit": unit,
        "maker": "",
        "note": "",
        "amount": amount or (price * qty if price and qty else 0),
    }


def _parse_structured_quote_text(text: str, vendor: str = ''):
    """PDF 추출 테이블이 없거나 깨지는 견적서 보조 파서.

    - 뉴스타자동화처럼 extract_tables()가 0개인 PDF
    - 에이엠테크처럼 금액이 `3 44,000`으로 분리되는 PDF
    - 강유시스템처럼 품목명 다음 줄에 규격이 이어지는 PDF
    """
    lines = [_normalize_split_amounts_in_line(l) for l in str(text or '').splitlines() if str(l or '').strip()]
    items = []
    last_item = None
    started = False
    for line in lines:
        n = _n(line)
        if any(k in n for k in ["합계", "소계", "납기일자", "유효일자", "결제조건", "비고", "공급자", "등록번호"]):
            if started and items:
                # 품목표 아래 조건/공급자 영역으로 내려가면 종료
                break
            continue
        parsed = _parse_structured_quote_line(line)
        if parsed:
            started = True
            items.append(parsed)
            last_item = parsed
            continue
        # 품목 직후 라인이 품번/규격만 있는 경우 이전 품목의 규격으로 병합
        cont = line.strip()
        if started and last_item and cont and not _has_num(cont.replace('-', '')) and _looks_part_no(cont):
            if not last_item.get('spec') or last_item.get('spec') == last_item.get('item_name'):
                last_item['spec'] = cont
        elif started and last_item and cont and _looks_part_no(cont) and not any(ch.isspace() for ch in cont):
            if not last_item.get('spec') or last_item.get('spec') == last_item.get('item_name'):
                last_item['spec'] = cont
    # 중복 제거
    out, seen = [], set()
    for it in items:
        key = (_n(it.get('item_name')), _n(it.get('spec')), int(it.get('quantity') or 0), int(it.get('amount') or 0))
        if key in seen:
            continue
        seen.add(key)
        out.append(it)
    return out


def _item_parse_score(items) -> int:
    score = 0
    for it in items or []:
        if str(it.get('item_name') or '').strip():
            score += 2
        if str(it.get('spec') or '').strip() and str(it.get('spec') or '').strip() != str(it.get('item_name') or '').strip():
            score += 2
        if int(float(it.get('quantity') or 0)) > 0:
            score += 1
        if int(float(it.get('unit_price') or 0)) > 1:
            score += 2
        if int(float(it.get('amount') or 0)) > 1:
            score += 2
        if int(float(it.get('unit_price') or 0)) <= 1 and int(float(it.get('amount') or 0)) <= 1:
            score -= 3
    return score


def _merge_or_prefer_text_items(table_items, text_items):
    table_items = list(table_items or [])
    text_items = list(text_items or [])
    if not text_items:
        return table_items
    if not table_items:
        return text_items
    table_bad = any(int(float(it.get('amount') or 0)) <= 1 and int(float(it.get('unit_price') or 0)) <= 1 for it in table_items)
    good_table_count = sum(1 for it in table_items if int(float(it.get('amount') or 0)) > 1 or int(float(it.get('unit_price') or 0)) > 1)
    if table_bad and len(text_items) >= max(1, good_table_count):
        return text_items
    if _item_parse_score(text_items) > _item_parse_score(table_items) + 2:
        return text_items
    # 표 인식 결과가 정상인 경우에는 텍스트 라인 보조 인식을 병합하지 않는다.
    # 병합하면 제이씨모션/강유시스템처럼 같은 품목이 다른 형태로 중복될 수 있다.
    return table_items


def _strip_auto_remarks(items):
    """자동인식 결과의 화면용 비고(제조사) 칸은 비운다.

    업체 구분은 vendor_name/vendor_index로 처리하므로 품목 행의 maker/note에
    업체명/납기/비고가 섞여 들어가지 않게 한다.
    """
    for it in items or []:
        it['maker'] = ''
        it['note'] = ''
    return items

def _fallback_text(text):
    items = _parse_multilingual_quote_text(text)
    if items:
        return items
    lines = [l.strip() for l in str(text or "").split("\n") if l.strip()]
    hi = -1
    for i, line in enumerate(lines):
        ll = _n(line)
        if ((("model" in ll or "modle" in ll or "품명" in ll) and "수량" in ll and ("단가" in ll or "금액" in ll))
                or _is_multilingual_quote_header(line)):
            hi = i
            break
    if hi < 0:
        return []
    for line in lines[hi+1:]:
        if _skip(line):
            continue
        tokens = [t.strip() for t in re.split(r"\s{2,}|\t", line) if t.strip()]
        if len(tokens) < 4:
            # pdfplumber 텍스트 추출이 공백 1개 기준일 때 보조
            tokens = [t.strip() for t in line.split() if t.strip()]
        nums = [int(re.sub(r"[^\d]", "", n)) for n in re.findall(r"[\d,]+", line) if re.sub(r"[^\d]", "", n)]
        if not tokens or not nums:
            continue
        # 품번 + MODEL + ... + 수량 + 단가 + 금액 형태 추정
        spec = next((t for t in tokens if _looks_part_no(t)), "")
        name = ""
        for t in tokens:
            if t == spec:
                continue
            if re.search(r"[A-Za-z가-힣]", t) and not _skip(t) and not _has_num(t):
                name = t
                break
        if not name:
            continue
        small = [n for n in nums if 1 <= n <= 9999]
        large = [n for n in nums if n > 100]
        qty = small[-1] if small else 1
        price = large[-2] if len(large) >= 2 else (large[0] if large else 0)
        amt = large[-1] if large else 0
        if price == 0 and amt and qty:
            price = amt // qty
        items.append({"item_name": name, "spec": spec, "unit_price": price, "quantity": qty, "unit": "EA", "maker": "", "amount": amt})
    return items



def _find_soffice_exe():
    """LibreOffice/soffice 실행 파일 탐색. 구형 .xls를 xlsx로 변환할 때만 사용한다."""
    for name in ("soffice", "libreoffice"):
        p = shutil.which(name)
        if p:
            return p
    if os.name == "nt":
        candidates = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ]
        for p in candidates:
            if os.path.exists(p):
                return p
    return ""


def _convert_xls_to_xlsx(path: str):
    """pandas/xlrd 없이 구형 .xls를 임시 xlsx로 변환한다.

    우선순위:
    1) Windows Excel COM: 서버 PC에 Excel이 있으면 가장 호환성이 좋음
    2) LibreOffice/soffice: 설치되어 있으면 headless 변환

    반환값: (xlsx_path, temp_dir, error_message)
    temp_dir은 호출자가 최종 정리한다.
    """
    path = os.path.abspath(path)
    tmp_dir = tempfile.mkdtemp(prefix="deverp_xls_")
    out_path = os.path.join(tmp_dir, os.path.splitext(os.path.basename(path))[0] + ".xlsx")
    errors = []

    if os.name == "nt":
        try:
            import pythoncom
            import win32com.client
            pythoncom.CoInitialize()
            excel = win32com.client.DispatchEx("Excel.Application")
            excel.Visible = False
            excel.DisplayAlerts = False
            wb = excel.Workbooks.Open(path)
            wb.SaveAs(out_path, FileFormat=51)  # xlOpenXMLWorkbook
            wb.Close(False)
            excel.Quit()
            try:
                pythoncom.CoUninitialize()
            except Exception:
                pass
            if os.path.exists(out_path) and os.path.getsize(out_path) > 0:
                return out_path, tmp_dir, ""
        except Exception as e:
            errors.append(f"Excel COM 변환 실패: {e}")
            try:
                excel.Quit()
            except Exception:
                pass
            try:
                pythoncom.CoUninitialize()
            except Exception:
                pass

    soffice = _find_soffice_exe()
    if soffice:
        try:
            cmd = [soffice, "--headless", "--convert-to", "xlsx", "--outdir", tmp_dir, path]
            cp = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, timeout=60)
            # LibreOffice가 파일명을 약간 바꿔 저장할 수 있어 가장 최신 xlsx를 찾는다.
            candidates = sorted(Path(tmp_dir).glob("*.xlsx"), key=lambda p: p.stat().st_mtime, reverse=True)
            if candidates and candidates[0].exists() and candidates[0].stat().st_size > 0:
                return str(candidates[0]), tmp_dir, ""
            errors.append(f"LibreOffice 변환 실패: {cp.stdout} {cp.stderr}".strip())
        except Exception as e:
            errors.append(f"LibreOffice 변환 실패: {e}")

    shutil.rmtree(tmp_dir, ignore_errors=True)
    return "", "", "; ".join(errors) or "구형 .xls 변환 도구를 찾지 못했습니다. Excel 또는 LibreOffice가 필요합니다."



def _looks_like_html_xls(raw: bytes) -> bool:
    head = raw[:8192].lower()
    return b"<html" in head or b"<table" in head or b"<tr" in head


def _decode_legacy_xls_html(raw: bytes) -> str:
    for enc in ("utf-8-sig", "cp949", "euc-kr", "utf-16", "latin1"):
        try:
            txt = raw.decode(enc)
            if "<" in txt and ">" in txt:
                return txt
        except Exception:
            pass
    return raw.decode("latin1", errors="ignore")


def _read_html_xls_rows(path: str):
    """Excel에서 .xls로 저장했지만 실제 내용은 HTML table인 견적서를 직접 읽는다."""
    try:
        raw = Path(path).read_bytes()
        if not _looks_like_html_xls(raw):
            return []
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(_decode_legacy_xls_html(raw), "html.parser")
        sheets = []
        for ti, table in enumerate(soup.find_all("table"), start=1):
            rows = []
            for tr in table.find_all("tr"):
                cells = []
                for cell in tr.find_all(["th", "td"]):
                    txt = cell.get_text(" ", strip=True)
                    cells.append(_s(txt))
                while cells and not cells[-1]:
                    cells.pop()
                if any(cells):
                    rows.append(cells)
            if rows:
                sheets.append((f"HTML_TABLE_{ti}", rows))
        return sheets
    except Exception:
        return []


def _xls_cell_text(cell, book) -> str:
    try:
        import xlrd
        if cell.ctype == xlrd.XL_CELL_EMPTY or cell.ctype == xlrd.XL_CELL_BLANK:
            return ""
        if cell.ctype == xlrd.XL_CELL_DATE:
            try:
                from datetime import datetime
                vals = xlrd.xldate_as_tuple(cell.value, book.datemode)
                if vals[3:] == (0, 0, 0):
                    return datetime(*vals).strftime("%Y-%m-%d")
                return datetime(*vals).strftime("%Y-%m-%d %H:%M")
            except Exception:
                return _s(cell.value)
        if cell.ctype == xlrd.XL_CELL_NUMBER:
            try:
                v = float(cell.value)
                return str(int(v)) if v.is_integer() else ("%f" % v).rstrip("0").rstrip(".")
            except Exception:
                return _s(cell.value)
        if cell.ctype == xlrd.XL_CELL_BOOLEAN:
            return "TRUE" if bool(cell.value) else "FALSE"
        return _s(cell.value)
    except Exception:
        return _s(getattr(cell, "value", ""))


def _read_xlrd_xls_rows(path: str):
    """Excel COM 없이 BIFF .xls 파일을 직접 읽는다. xlrd가 없으면 빈 목록을 반환한다."""
    try:
        import xlrd
        book = xlrd.open_workbook(path, on_demand=True)
        sheets = []
        for sh in book.sheets():
            rows = []
            max_r = min(sh.nrows, 300)
            max_c = min(sh.ncols, 80)
            for r in range(max_r):
                row = []
                for c in range(max_c):
                    row.append(_xls_cell_text(sh.cell(r, c), book))
                while row and not row[-1]:
                    row.pop()
                if any(row):
                    rows.append(row)
            if rows:
                sheets.append((str(sh.name), rows))
        try:
            book.release_resources()
        except Exception:
            pass
        return sheets
    except Exception:
        return []


def _read_legacy_xls_rows_direct(path: str):
    """구형 .xls를 Excel COM/LibreOffice 없이 먼저 직접 읽는다.

    처리 순서:
    1) HTML table 형식 .xls
    2) BIFF .xls + xlrd
    """
    sheets = _read_html_xls_rows(path)
    if sheets:
        return sheets, ""
    sheets = _read_xlrd_xls_rows(path)
    if sheets:
        return sheets, ""
    return [], "직접 .xls 읽기 실패 또는 xlrd 미설치"

def _parse_pdf(path):
    import pdfplumber
    items, vendor, full_text = [], "", ""
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            full_text += text + "\n"
            for tbl in (page.extract_tables() or []):
                found = _parse_table(tbl)
                if not found:
                    found = _parse_loose_rows(tbl)
                if found:
                    items.extend(found)
    if not vendor:
        vendor = _vendor_from_text(full_text)
    # PDF 테이블 인식 결과가 없거나, 금액/규격이 깨진 경우 텍스트 라인 기반 보조 인식으로 보정한다.
    text_items = _parse_structured_quote_text(full_text, vendor)
    if text_items:
        items = _merge_or_prefer_text_items(items, text_items)
    if not items:
        items = _parse_multilingual_quote_text(full_text, vendor)
    if not items:
        items = _fallback_text(full_text)
    if not items:
        # 텍스트 라인 기반 마지막 보조 인식
        line_rows = [[x] for x in full_text.splitlines()]
        items = _parse_loose_rows(line_rows) or _parse_amount_anchored_rows(line_rows)
    items = _normalize_and_filter_items(items)
    _enrich_items_hidden_fields(items, full_text + " " + os.path.basename(path))
    _strip_auto_remarks(items)
    return {"items": items, "vendor_name": vendor}

def _parse_xlsx(path):
    ext = os.path.splitext(path)[1].lower()
    items, vendor = [], ""
    tmp_dir = ""
    try:
        all_rows = []
        sheet_rows_list = []
        context_parts = [os.path.basename(path)]
        read_path = path

        if ext == ".xls":
            # Excel COM이 없는 서버/SYSTEM 부팅 실행 환경에서도 .xls가 인식되도록
            # HTML-xls 또는 BIFF-xls(xlrd)를 먼저 직접 읽는다.
            direct_sheets, _direct_err = _read_legacy_xls_rows_direct(path)
            if direct_sheets:
                for sheet_name, rows in direct_sheets:
                    context_parts.append(str(sheet_name))
                    rows = [[_s(v) for v in row] for row in rows if any(_s(v) for v in row)]
                    if rows:
                        sheet_rows_list.append(rows)
                        all_rows.extend(rows)
                        context_parts.append(" ".join(" ".join(r) for r in rows[:20]))
                read_path = ""
            else:
                converted, tmp_dir, conv_err = _convert_xls_to_xlsx(path)
                if converted:
                    read_path = converted
                else:
                    return {"items": [], "vendor_name": "", "error": conv_err or _direct_err}

        if read_path:
            import openpyxl
            wb = openpyxl.load_workbook(read_path, read_only=False, data_only=True)
            for ws in wb.worksheets:
                context_parts.append(str(ws.title))
                # 바이러스/매크로 안내 시트처럼 견적과 무관한 시트는 자연스럽게 무시된다.
                merged = {}
                try:
                    ranges = ws.merged_cells.ranges
                except Exception:
                    ranges = []
                for mr in ranges:
                    val = ws.cell(mr.min_row, mr.min_col).value
                    for rr in range(mr.min_row, mr.max_row + 1):
                        for cc in range(mr.min_col, mr.max_col + 1):
                            merged[(rr, cc)] = val
                rows = []
                max_r = min(ws.max_row or 0, 300)
                max_c = min(ws.max_column or 0, 80)
                for rr in range(1, max_r + 1):
                    row = []
                    for cc in range(1, max_c + 1):
                        v = merged.get((rr, cc), ws.cell(rr, cc).value)
                        row.append(_s(v))
                    while row and not row[-1]:
                        row.pop()
                    if any(row):
                        rows.append(row)
                if rows:
                    sheet_rows_list.append(rows)
                    all_rows.extend(rows)
                    context_parts.append(" ".join(" ".join(r) for r in rows[:20]))

        vendor = _vendor_from_rows(all_rows)
        for rows in sheet_rows_list:
            found = _parse_table(rows)
            if not found:
                found = _parse_numeric_quote_rows(rows)
            if not found:
                found = _parse_loose_rows(rows)
            if not found:
                found = _parse_amount_anchored_rows(rows)
            if found:
                items.extend(found)
        if not items:
            items = (_parse_table(all_rows) or _parse_numeric_quote_rows(all_rows) or _parse_loose_rows(all_rows) or _parse_amount_anchored_rows(all_rows))
    except Exception as e:
        return {"items": [], "vendor_name": "", "error": str(e)}
    finally:
        if tmp_dir:
            shutil.rmtree(tmp_dir, ignore_errors=True)
    out, seen = [], set()
    for it in items:
        try:
            q = int(float(it.get('quantity') or 0))
        except Exception:
            q = 0
        try:
            a = int(float(it.get('amount') or 0))
        except Exception:
            a = 0
        key = (_n(it.get('item_name')), _n(it.get('spec')), q, a)
        if key in seen:
            continue
        seen.add(key)
        out.append(it)
    out = _normalize_and_filter_items(out)
    _enrich_items_hidden_fields(out, " ".join(context_parts) + " " + os.path.basename(path))
    return {"items": out, "vendor_name": vendor}

def _parse_docx(path):
    from docx import Document
    doc = Document(path)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    vendor = _vendor_from_text(full_text)
    items = []
    for tbl in doc.tables:
        rows = [[c.text.strip() for c in row.cells] for row in tbl.rows]
        found = _parse_table(rows)
        if not found:
            found = _parse_loose_rows(rows)
        if found:
            items.extend(found)
    if not items:
        items = _fallback_text(full_text)
    if not items:
        line_rows = [[x] for x in full_text.splitlines()]
        items = _parse_loose_rows(line_rows) or _parse_amount_anchored_rows(line_rows)
    items = _normalize_and_filter_items(items)
    table_text = " ".join(" ".join(c.text.strip() for c in row.cells) for tbl in doc.tables for row in tbl.rows)
    _enrich_items_hidden_fields(items, full_text + " " + table_text + " " + os.path.basename(path))
    return {"items": items, "vendor_name": vendor}

def parse_estimate(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    try:
        if ext == ".pdf":
            result = _parse_pdf(file_path)
        elif ext in [".xlsx", ".xlsm", ".xls"]:
            result = _parse_xlsx(file_path)
        elif ext in [".docx", ".doc"]:
            result = _parse_docx(file_path)
        else:
            result = {"items": [], "vendor_name": "", "error": f"지원 안 함: {ext}"}
        # 파서별 보정이 누락되어도 파일명 기준으로 한 번 더 보정한다.
        result["items"] = _normalize_and_filter_items(result.get("items") or [])
        _enrich_items_hidden_fields(result.get("items") or [], os.path.basename(file_path))
        _strip_auto_remarks(result.get("items") or [])
        return result
    except Exception as e:
        return {"items": [], "vendor_name": "", "error": str(e)}

def parse_multiple(file_paths):
    all_items, vendors, errors, files = [], [], [], []
    for fp in file_paths:
        r = parse_estimate(fp)
        vendor = r.get("vendor_name") or os.path.splitext(os.path.basename(fp))[0]
        if vendor and vendor not in vendors:
            vendors.append(vendor)
        f_items = []
        for item in r.get("items", []):
            item = dict(item)
            item["vendor_name"] = vendor
            f_items.append(item)
            all_items.append(item)
        files.append({"file": fp, "vendor_name": vendor, "items": f_items, "error": r.get("error", "")})
        if r.get("error"):
            errors.append(f"{os.path.basename(fp)}: {r['error']}")
        elif not f_items:
            errors.append(f"{os.path.basename(fp)}: 인식된 품목 없음")
    return {"items": all_items, "vendors": vendors, "errors": errors, "files": files}
